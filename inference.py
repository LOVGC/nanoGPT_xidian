from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable

import httpx
import torch
import tiktoken
from transformers.generation.logits_process import (
    TemperatureLogitsWarper,
    TopKLogitsWarper,
)

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma

from PyPDF2 import PdfReader
import pandas as pd
import docx

import model


DEFAULT_DOCS_DIR = Path("data") / "knowledge_base"
DEFAULT_INDEX_DIR = Path("rag_index")
DEFAULT_EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
DEFAULT_OLLAMA_BASE = "http://localhost:11434"
DEFAULT_OLLAMA_MODEL = "qwen3:4b"
SYSTEM_PROMPT = (
    "You are a RAG assistant. Use the provided context to answer the question. "
    'If the answer is not in the context, say "I don\'t know."'
)


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: Path) -> Iterable[Document]:
    reader = PdfReader(str(path))
    for page_idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            yield Document(
                page_content=text, metadata={"source": str(path), "page": page_idx}
            )


def load_docx(path: Path) -> Iterable[Document]:
    document = docx.Document(str(path))
    text = "\n".join([p.text for p in document.paragraphs if p.text])
    if text.strip():
        yield Document(page_content=text, metadata={"source": str(path)})


def load_tabular(path: Path) -> Iterable[Document]:
    if path.suffix.lower() == ".csv":
        df = pd.read_csv(path)
    else:
        df = pd.read_excel(path)
    text = df.to_string(index=False)
    if text.strip():
        yield Document(page_content=text, metadata={"source": str(path)})


def load_documents(docs_dir: Path) -> list[Document]:
    documents: list[Document] = []
    if not docs_dir.exists():
        return documents

    for path in docs_dir.rglob("*"):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            text = read_text_file(path)
            if text.strip():
                documents.append(
                    Document(page_content=text, metadata={"source": str(path)})
                )
        elif suffix == ".pdf":
            documents.extend(list(load_pdf(path)))
        elif suffix in {".docx", ".doc"}:
            documents.extend(list(load_docx(path)))
        elif suffix in {".csv", ".xlsx", ".xls"}:
            documents.extend(list(load_tabular(path)))
    return documents


def build_vectorstore(
    docs_dir: Path,
    index_dir: Path,
    rebuild: bool,
    embed_model: str,
) -> Chroma:
    embedding = HuggingFaceEmbeddings(model_name=embed_model)

    index_dir.mkdir(parents=True, exist_ok=True)
    has_index = any(index_dir.iterdir())

    if not rebuild and has_index:
        return Chroma(persist_directory=str(index_dir), embedding_function=embedding)

    documents = load_documents(docs_dir)
    if not documents:
        raise RuntimeError(f"No documents found in {docs_dir}")

    splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=80)
    chunks = splitter.split_documents(documents)

    vectordb = Chroma.from_documents(
        documents=chunks,
        embedding=embedding,
        persist_directory=str(index_dir),
    )
    try:
        vectordb.persist()
    except AttributeError:
        pass
    return vectordb


def build_prompt(
    instruction: str,
    question: str,
    context: str,
    enc: tiktoken.Encoding,
    max_input_tokens: int,
) -> str:
    base_no_question = (
        "### Instruction:\n"
        + instruction
        + "\n### Input:\nContext:\n"
        + "\n\nQuestion:\n"
        + "\n### Response:\n"
    )

    question_tokens = enc.encode(question)
    base_tokens = enc.encode(base_no_question)
    available_for_question = max_input_tokens - len(base_tokens)
    if available_for_question < 0:
        available_for_question = 0
    if len(question_tokens) > available_for_question:
        question_tokens = question_tokens[-available_for_question:]
    question_text = enc.decode(question_tokens)

    prompt_without_context = (
        "### Instruction:\n"
        + instruction
        + "\n### Input:\nContext:\n"
        + "\n\nQuestion:\n"
        + question_text
        + "\n### Response:\n"
    )
    prompt_tokens = enc.encode(prompt_without_context)
    available_for_context = max_input_tokens - len(prompt_tokens)
    if available_for_context < 0:
        available_for_context = 0
    context_tokens = enc.encode(context)
    context_tokens = context_tokens[:available_for_context]
    context_text = enc.decode(context_tokens)

    return (
        "### Instruction:\n"
        + instruction
        + "\n### Input:\nContext:\n"
        + context_text
        + "\n\nQuestion:\n"
        + question_text
        + "\n### Response:\n"
    )


def build_user_message(
    question: str,
    context: str,
    enc: tiktoken.Encoding,
    max_input_tokens: int,
) -> str:
    base_no_context = "Context:\n\nQuestion:\n" + question
    base_tokens = enc.encode(base_no_context)
    available_for_context = max_input_tokens - len(base_tokens)
    if available_for_context < 0:
        available_for_context = 0
    context_tokens = enc.encode(context)
    context_tokens = context_tokens[:available_for_context]
    context_text = enc.decode(context_tokens)
    return "Context:\n" + context_text + "\n\nQuestion:\n" + question


def retrieve_context(vectordb: Chroma, question: str, top_k: int) -> str:
    docs = vectordb.similarity_search(question, k=top_k)
    if not docs:
        return ""
    return "\n\n".join([doc.page_content for doc in docs])


def load_local_model(checkpoint_path: Path, device: torch.device):
    checkpoint = torch.load(checkpoint_path, map_location=device)
    model_args = model.ModelArgs(**checkpoint["model_args"])
    model_instance = model.Transformer(model_args).to(device)
    model_instance.load_state_dict(checkpoint["model_state_dict"])
    model_instance.eval()
    return model_instance, model_args


def generate_local(
    model_instance: torch.nn.Module,
    enc: tiktoken.Encoding,
    prompt: str,
    max_new_tokens: int,
    temperature: float,
    top_k: int,
    block_size: int,
) -> str:
    input_ids = torch.tensor(
        enc.encode(prompt),
        dtype=torch.long,
        device=next(model_instance.parameters()).device,
    ).unsqueeze(0)
    warpers = []
    if temperature and temperature != 1.0:
        warpers.append(TemperatureLogitsWarper(temperature))
    if top_k and top_k > 0:
        warpers.append(TopKLogitsWarper(top_k=top_k))

    eot_id = enc.eot_token
    for _ in range(max_new_tokens):
        input_cond = input_ids[:, -block_size:]
        logits = model_instance(input_cond)
        scores = logits[:, -1, :]
        for warper in warpers:
            scores = warper(input_ids, scores)
        probs = torch.softmax(scores, dim=-1)
        next_token = torch.multinomial(probs, num_samples=1)
        input_ids = torch.cat([input_ids, next_token], dim=1)
        if next_token.item() == eot_id:
            break
    output_ids = input_ids[0, len(enc.encode(prompt)) :].tolist()
    if eot_id in output_ids:
        output_ids = output_ids[: output_ids.index(eot_id)]
    return enc.decode(output_ids).strip()


def call_ollama_generate(
    base_url: str,
    model_name: str,
    prompt: str,
    temperature: float,
    top_k: int,
    top_p: float,
) -> str:
    url = base_url.rstrip("/") + "/api/generate"
    payload = {
        "model": model_name,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": temperature,
            "top_k": top_k,
            "top_p": top_p,
        },
    }
    with httpx.Client(timeout=60.0) as client:
        response = client.post(url, json=payload)
        response.raise_for_status()
        data = response.json()
    return data.get("response", "").strip()


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="RAG inference with local or Ollama backend"
    )
    parser.add_argument("--backend", choices=["local", "ollama"], default="local")
    parser.add_argument("--checkpoint", type=str, default="model_trained/sft_best.pt")
    parser.add_argument("--docs_dir", type=str, default=str(DEFAULT_DOCS_DIR))
    parser.add_argument("--index_dir", type=str, default=str(DEFAULT_INDEX_DIR))
    parser.add_argument("--rebuild", action="store_true")
    parser.add_argument("--top_k", type=int, default=3)
    parser.add_argument("--embed_model", type=str, default=DEFAULT_EMBED_MODEL)
    parser.add_argument("--max_new_tokens", type=int, default=128)
    parser.add_argument("--temperature", type=float, default=0.7)
    parser.add_argument("--top_k_sample", type=int, default=40)
    parser.add_argument("--ollama_base_url", type=str, default=DEFAULT_OLLAMA_BASE)
    parser.add_argument("--ollama_model", type=str, default=DEFAULT_OLLAMA_MODEL)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    docs_dir = Path(args.docs_dir)
    index_dir = Path(args.index_dir)

    vectordb = build_vectorstore(
        docs_dir=docs_dir,
        index_dir=index_dir,
        rebuild=args.rebuild,
        embed_model=args.embed_model,
    )

    instruction = (
        "Answer the question using the context. "
        'If the answer is not in the context, say "I don\'t know."'
    )

    if args.backend == "local":
        device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        model_instance, model_args = load_local_model(Path(args.checkpoint), device)
        enc = tiktoken.get_encoding("gpt2")
        max_input_tokens = model_args.max_seq_len
    else:
        model_instance = None
        model_args = None
        enc = tiktoken.get_encoding("gpt2")
        max_input_tokens = 512

    print(f"Backend: {args.backend}")
    print(f"Docs dir: {docs_dir}")
    print("Type 'exit' to quit.")

    while True:
        question = input("> ").strip()
        if question.lower() in {"exit", "quit", ""}:
            break

        context = retrieve_context(vectordb, question, args.top_k)
        if args.backend == "local":
            prompt = build_prompt(
                instruction=instruction,
                question=question,
                context=context,
                enc=enc,
                max_input_tokens=max_input_tokens,
            )
            assert model_instance is not None
            response = generate_local(
                model_instance,
                enc,
                prompt,
                max_new_tokens=args.max_new_tokens,
                temperature=args.temperature,
                top_k=args.top_k_sample,
                block_size=max_input_tokens,
            )
        else:
            user_message = build_user_message(
                question=question,
                context=context,
                enc=enc,
                max_input_tokens=max_input_tokens,
            )
            full_prompt = f"System: {SYSTEM_PROMPT}\n\n{user_message}"
            response = call_ollama_generate(
                base_url=args.ollama_base_url,
                model_name=args.ollama_model,
                prompt=full_prompt,
                temperature=args.temperature,
                top_k=args.top_k_sample,
                top_p=0.9,
            )

        print(response)


if __name__ == "__main__":
    main()
